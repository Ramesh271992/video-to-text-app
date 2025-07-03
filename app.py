import streamlit as st
from docx import Document
import os
import whisper

# Force whisper to use this ffmpeg path
os.environ["PATH"] = r"C:\ffmpeg\ffmpeg-7.1.1-essentials_build\bin;" + os.environ["PATH"]

st.set_page_config(page_title="🎥 Video to Text AI", layout="centered")
st.title("🎙️ Video to Text (Telugu/English) Converter")

uploaded_file = st.file_uploader("📤 Upload your video/audio file", type=["mp4", "mp3", "wav", "m4a", "mov"])

if uploaded_file is not None:
    with open("input_file", "wb") as f:
        f.write(uploaded_file.read())
    st.success("✅ File uploaded successfully!")

    if st.button("🧠 Transcribe"):
        st.info("⏳ Transcribing... Please wait.")
        model = whisper.load_model("large")
        result = model.transcribe("input_file")
        transcription = result["text"]

        # Display in app
        st.text_area("📄 Transcript:", transcription, height=300)

        # Save as .docx
        doc = Document()
        doc.add_heading("Video Transcript", level=1)
        doc.add_paragraph(transcription)
        doc.save("transcript.docx")
        with open("transcript.docx", "rb") as file:
            st.download_button("⬇️ Download Word Document", file, file_name="transcript.docx")

        # Save as .txt
        with open("transcript.txt", "w", encoding="utf-8") as file:
            file.write(transcription)
        with open("transcript.txt", "rb") as file:
            st.download_button("⬇️ Download TXT File", file, file_name="transcript.txt")
